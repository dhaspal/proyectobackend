# -*- coding: utf-8 -*-
"""Genera los 4 informes markdown (prompt + resultado) y el DOCX consolidado."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
PROJ = ROOT.parent
PROMPT_PATH = ROOT / "Prompt_informe_de_vulnerabilidades_en_codigo.txt"
DOCX_OUT = PROJ / "docs" / "informe-busqueda-vulnerabilidades-codigo.docx"


def load_prompt() -> str:
    return PROMPT_PATH.read_text(encoding="utf-8")


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


# --- Textos de análisis por carpeta (cumplen estructura del prompt; CVE solo donde aplica) ---

RESULT_LOGIN = r"""
## 1. Escaneo y enumeración

- **Tipo de aplicación:** API REST Node.js (Express), autenticación JWT, bcrypt.
- **Archivos analizados (carpeta Login):** `auth.routes.js`, `auth.controller.js`, `auth.validators.js`, `tokens.js`, `password.js`, `middlewares/auth.js`.
- **Metodología:** revisión estática de código; sin pentesting activo.

## Resumen ejecutivo

El módulo de autenticación implementa login/registro con validación Zod y JWT; se observan riesgos de **política de contraseña débil**, **registro público con rol `MECHANIC`**, **TTL de JWT largo por defecto**, y **ausencia de rate-limit específico** en endpoints de autenticación (solo limitador global en `/api`).

## Tabla de salida

| Vulnerabilidad — NOMBRE (CWE / CVE) | Descripción | Enlaces CWE, CVE, MITRE, NIST, CVSS | Vector NIST (ejemplo) | Código vulnerable | Impacto | Código corregido (propuesto) | Mitigación | Verificación | ISO / OWASP |
|---|---|---|---|---|---|---|---|---|---|
| Política de contraseña insuficiente (CWE-521 / sin CVE) | Mínimo 6 caracteres. | https://cwe.mitre.org/data/definitions/521.html | N/A | `password: z.string().min(6)` en `auth.validators.js` | Favorece fuerza bruta y credential stuffing. | `min(12)` + reglas de complejidad | Política corporativa + validación | Tests Zod | A.8.5 / A07:2021 |
| Elevación lógica en registro (CWE-269 / sin CVE) | Rol `MECHANIC` seleccionable en registro público. | https://cwe.mitre.org/data/definitions/269.html | N/A | `role: z.enum([CLIENT, MECHANIC])` | Cuentas de taller sin control de confianza. | Forzar `CLIENT` en registro público | Alta de mecánico solo por admin | Test registro | A.9.2 / A01:2021 |
| Sesión JWT larga (CWE-613 / sin CVE) | Default `7d`. | https://cwe.mitre.org/data/definitions/613.html | N/A | `getEnvOptional("JWT_EXPIRES_IN", "7d")` | Ventana de abuso si roban token. | TTL corto + refresh | Config + revocación | Test expiración | A.9.4 / A07:2021 |
| Ausencia de límites dedicados en auth (CWE-307 estilo / sin CVE) | `/auth/login` comparte limitador global. | https://cwe.mitre.org/data/definitions/307.html | N/A | Rutas montadas bajo `/api` con `apiLimiter` único | Posible abuso de credenciales a mayor tasa relativa. | `rateLimit` específico en `auth.routes` | Controles por IP/usuario | Tests de 429 en login | A.8.32 / A07:2021 |

## [#1] Vulnerabilidad — Política de contraseña débil (CWE-521 / sin CVE público)

- **CWE:** https://cwe.mitre.org/data/definitions/521.html
- **CVE:** no aplica al código propio.
- **CVSS publicado en NVD:** no aplica al hallazgo de código (no hay CVE).

**Evidencia (vulnerable):**
```js
password: z.string().min(6).max(200),
```

**Corrección (propuesta):**
```js
password: z.string().min(12).max(200).regex(/^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[^A-Za-z0-9]).+$/)
```

**Nota estudiantes:** la longitud mínima es el control más rentable frente a ataques offline.

## [#2] Vulnerabilidad — Registro con rol mecánico (CWE-269 / sin CVE público)

**Evidencia (vulnerable):**
```js
role: z.enum([ROLES.CLIENT, ROLES.MECHANIC]).default(ROLES.CLIENT),
```

**Corrección (propuesta):** `role: z.literal(ROLES.CLIENT)` y en controlador ignorar rol del body.

## Conclusiones y recomendaciones

Endurecer políticas de identidad, eliminar privilegio en registro público, acortar JWT y añadir rate-limit específico a login/register.

## Recursos y referencias

- CWE: https://cwe.mitre.org/
- OWASP Authentication: https://owasp.org/www-project-top-ten/
- CVSS v3.1 especificación: https://www.first.org/cvss/v3.1/specification-document
"""

RESULT_FRONTS = r"""
## 1. Escaneo y enumeración

- **Tipo de aplicación:** API REST Node.js + Mongoose (dominio taller: citas, OT, vehículos, finanzas, etc.).
- **Archivos analizados:** árbol `Fronts/src` (rutas, controladores, modelos, validadores, middlewares `errorHandler`, `notFound`, `rateLimit`, servicios).
- **Metodología:** revisión estática; sin pentesting activo.

## Resumen ejecutivo

La superficie expone numerosos recursos con controles por rol en muchos endpoints; persisten riesgos típicos de **validación parcial en casos administrativos**, **filtración en errores** (manejador global compartido con Config), y **lógica de negocio** que debe mantenerse alineada con autorización (IDOR ya mitigado en varios controladores).

## Tabla de salida (extracto representativo)

| Vulnerabilidad — NOMBRE (CWE / CVE) | Descripción | Enlaces | Vector | Evidencia | Impacto | Corrección propuesta | Mitigación | Verificación | ISO / OWASP |
|---|---|---|---|---|---|---|---|---|---|
| Validación incompleta de `mechanicId` en reclamo admin (CWE-20 / sin CVE) | `req.body.mechanicId` sin Zod en `claimWorkOrder`. | https://cwe.mitre.org/data/definitions/020.html | N/A | `workOrders.controller.js` | Datos inconsistentes / fallos de asignación | Esquema Zod + comprobar usuario mecánico | Validación estricta | Tests API | A.14.1 / API3:2023 |
| Divulgación vía errores (CWE-209 / sin CVE) | Respuestas con metadatos sensibles si se reutiliza handler global. | https://cwe.mitre.org/data/definitions/209.html | N/A | `middlewares/errorHandler.js` | Reconocimiento / enumeración | Sanitizar en producción | Errores genéricos | Tests | A.8.15 / A04:2021 |

## [#1] Vulnerabilidad — Validación de entrada incompleta en reclamo de OT (CWE-20 / sin CVE público)

**Evidencia (vulnerable):**
```js
wo.mechanic = req.user.role === ROLES.ADMIN && req.body?.mechanicId ? req.body.mechanicId : req.user.sub;
```

**Corrección (propuesta):** parsear `req.body` con Zod (`mechanicId` como ObjectId) y verificar `User.exists({ _id, role: MECHANIC })`.

## Conclusiones y recomendaciones

Mantener validación Zod en **todos** los endpoints que acepten JSON; revisar periodicamente controladores con ramas `ADMIN`.

## Recursos y referencias

- OWASP API Top 10: https://owasp.org/www-project-api-security/
- CWE: https://cwe.mitre.org/
"""

RESULT_CONEXION = r"""
## 1. Escaneo y enumeración

- **Tipo:** conexión MongoDB mediante Mongoose; script adicional con driver nativo para migración.
- **Archivos:** `db.js`, `index.js` (bootstrap), `migrate-local-to-atlas.js`.
- **Metodología:** revisión estática; sin pruebas contra bases reales.

## Resumen ejecutivo

La URI se obtiene de variables de entorno (**bien**); `strictQuery` activado (**bien**). Riesgos: **autoIndex en no producción** reduce diferencias entorno prod/test; script de migración **borra destino** (`deleteMany`) — riesgo operativo elevado si se ejecuta contra producción sin controles.

## Tabla de salida

| Vulnerabilidad — NOMBRE (CWE / CVE) | Descripción | Enlaces | Vector | Evidencia | Impacto | Corrección | Mitigación | Verificación | ISO / OWASP |
|---|---|---|---|---|---|---|---|---|---|
| Operación destructiva en migración (CWE-432 / sin CVE) | `deleteMany({})` antes de copiar. | https://cwe.mitre.org/data/definitions/432.html | N/A | `migrate-local-to-atlas.js` | Pérdida de datos en Atlas | Modo dry-run / confirmación explícita | Procedimientos + backups | Revisiones de script | A.12.1 |
| Cadena de conexión en env (CWE-798 estilo / sin CVE propio) | Credenciales vía `MONGO_URI`/`ATLAS_URI` — correcto si `.env` no se versiona. | https://cwe.mitre.org/data/definitions/798.html | N/A | `getEnv("MONGO_URI")` | Exposición si `.env` filtrado | Secret manager | No commitear secretos | Gitleaks | A.8.2 |

## [#1] Vulnerabilidad — Riesgo operativo por migración destructiva (clasificación orientativa CWE-432 / sin CVE público)

**Evidencia:**
```js
await dst.deleteMany({});
```

**Corrección (propuesta):** flag `CONFIRM_DESTRUCTIVE=1`, backup obligatorio, o escritura incremental sin borrado total por defecto.

## Conclusiones y recomendaciones

Proteger scripts de migración con confirmación explícita y entornos separados.

## Recursos

- MongoDB security checklist (documentación oficial MongoDB Inc.)
- CWE: https://cwe.mitre.org/
"""

RESULT_CONFIG = r"""
## 1. Escaneo y enumeración

- **Stack:** Node.js, Express 5, dependencias listadas en `package.json`, configuración en `app.js`, variables con `env.js`.
- **Metodología:** revisión estática + correlación con `npm audit` (dependencias).

## Resumen ejecutivo

Se identifica **mala configuración de CORS** cuando `FRONTEND_ORIGINS` está vacío (permite cualquier origen con `Origin` presente según la lógica actual). En dependencias directas/transitivas aparecen CVE publicados (ver NVD).

## Tabla de salida

| Vulnerabilidad — NOMBRE (CWE / CVE) | Descripción | Enlaces | Vector NIST | Evidencia | Impacto | Corrección | Mitigación | Verificación | ISO / OWASP |
|---|---|---|---|---|---|---|---|---|---|
| CORS permisivo por defecto (CWE-942 / sin CVE) | Lista vacía ⇒ `cb(null, true)`. | https://cwe.mitre.org/data/definitions/942.html | N/A | `app.js` | Superficie ampliada | Fail-closed en prod | `FRONTEND_ORIGINS` obligatorio | Tests CORS | A.8.28 / API8:2023 |
| CVE-2026-30827 express-rate-limit | Colisión bucket IPv4-mapped IPv6. | CWE-770 https://cwe.mitre.org/data/definitions/770.html · MITRE https://cve.mitre.org/cgi-bin/cvename.cgi?name=CVE-2026-30827 · NIST https://nvd.nist.gov/vuln/detail/CVE-2026-30827 · GHSA https://github.com/advisories/GHSA-46wh-pxpv-q5gq | CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:N/A:H (7.5) | `package.json` ^8.2.1 | DoS lógico rate limit | Upgrade >=8.2.2 | `npm update` | Pruebas dual-stack | A.12.6 / A06:2021 |
| CVE-2026-4926 path-to-regexp | ReDoS grupos opcionales. | CWE-1333 · MITRE · NIST https://nvd.nist.gov/vuln/detail/CVE-2026-4926 · GHSA-j3q9-mxjg-w52f | CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:N/A:H | Transitiva Express | DoS CPU | Actualizar cadena Express | `npm audit fix` | Tests de rutas | A.12.6 |
| CVE-2026-4923 path-to-regexp | ReDoS comodines. | NIST https://nvd.nist.gov/vuln/detail/CVE-2026-4923 | CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:N/I:N/A:H (5.9) | Transitiva | DoS | Upgrade path-to-regexp ≥8.4.0 | Idem | Idem | A.12.6 |

**CVSS documentación (no usar calculadora como sustituto del registro CVE):** https://www.first.org/cvss/v3.1/specification-document

## [#1] Vulnerabilidad — CORS permisivo (CWE-942 / sin CVE público)

**Evidencia:**
```js
if (allowedOrigins.length === 0) return cb(null, true);
```

**Corrección:** rechazar arranque en producción si la lista está vacía.

## Conclusiones

Actualizar dependencias con CVE; fijar CORS estricto; ejecutar SCA en CI.

## Recursos

- NVD: https://nvd.nist.gov/
- FIRST CVSS v3.1: https://www.first.org/cvss/v3.1/specification-document
"""


def write_md(name: str, body: str) -> Path:
    prompt = load_prompt()
    out = ROOT / name
    out.write_text(
        "# Informe generado (prompt + resultado)\n\n"
        "## Prompt compartido (texto íntegro)\n\n"
        "```\n"
        + prompt
        + "\n```\n\n"
        "## Resultado del análisis\n"
        + body,
        encoding="utf-8",
    )
    return out


def build_docx():
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Informe de búsqueda de vulnerabilidades en código")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph(
        "Proyecto: proyectogradobackend — API Node.js/Express/Mongoose."
    )
    doc.add_paragraph(
        "Metodología: revisión estática de código y de dependencias; sin explotación activa."
    )

    add_heading(doc, "Resumen ejecutivo consolidado", 1)
    add_para(
        doc,
        "El sistema presenta debilidades en autenticación/autorización de políticas (contraseñas, roles en registro, TTL JWT), "
        "configuración de CORS, validación puntual en flujos administrativos, y vulnerabilidades conocidas en dependencias "
        "(express-rate-limit CVE-2026-30827; path-to-regexp CVE-2026-4926 y CVE-2026-4923; entre otras transitivas reportadas por npm audit). "
        "Se recomienda remediación priorizada y repetición del análisis tras parches.",
    )

    add_heading(doc, "Referencias NIST / MITRE / CVE citados", 2)
    add_para(
        doc,
        "CVE-2026-30827 (NIST): https://nvd.nist.gov/vuln/detail/CVE-2026-30827\n"
        "CVE-2026-4926 (NIST): https://nvd.nist.gov/vuln/detail/CVE-2026-4926\n"
        "CVE-2026-4923 (NIST): https://nvd.nist.gov/vuln/detail/CVE-2026-4923\n"
        "CVSS v3.1 especificación FIRST: https://www.first.org/cvss/v3.1/specification-document",
    )

    def append_markdownish(doc, title, text):
        add_heading(doc, title, 1)
        for line in text.strip().splitlines():
            s = line.strip()
            if not s:
                continue
            if s.startswith("#### "):
                add_heading(doc, s[5:], 4)
            elif s.startswith("### "):
                add_heading(doc, s[4:], 3)
            elif s.startswith("## "):
                add_heading(doc, s[3:], 2)
            elif s.startswith("# "):
                add_heading(doc, s[2:], 1)
            elif s.startswith("|") and "---" not in s:
                add_para(doc, s)
            elif s.startswith("```"):
                continue
            else:
                add_para(doc, s)

    for title, body in [
        ("Carpeta Login — resultado", RESULT_LOGIN),
        ("Carpeta Fronts — resultado", RESULT_FRONTS),
        ("Carpeta Conexion — resultado", RESULT_CONEXION),
        ("Carpeta Configuracion — resultado", RESULT_CONFIG),
    ]:
        append_markdownish(doc, title, body)

    doc.save(str(DOCX_OUT))
    return DOCX_OUT


def main():
    write_md("informe_resultado_01_Login.md", RESULT_LOGIN)
    write_md("informe_resultado_02_Fronts.md", RESULT_FRONTS)
    write_md("informe_resultado_03_Conexion.md", RESULT_CONEXION)
    write_md("informe_resultado_04_Configuracion.md", RESULT_CONFIG)
    out = build_docx()
    print("Generado:", out)


if __name__ == "__main__":
    main()
